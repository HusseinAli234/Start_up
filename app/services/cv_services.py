import asyncio
import pdfplumber
import io # For reading bytes
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from fastapi import HTTPException
from docx import Document
from google.cloud import storage # Import GCS client
import logging

from app.models.employers import JobPosting
from app.schemas.vacancy_schema import SkillSchema
from app.ai.analyzer import analyze_resume,analyze_resume_chatgpt
from fastapi.responses import JSONResponse

logger = logging.getLogger(__name__)
class CVService:
    def __init__(self, db: AsyncSession):
        self.db = db

    def _read_pdf(self, file_path: str) -> str:
        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(layout=True)
                if page_text:
                    text.append(page_text)           
        return "\n".join(text)
    def _read_docx_content(self, docx_content: bytes) -> str:
        try:
            doc = Document(io.BytesIO(docx_content))
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())

            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error reading DOCX content: {e}")
            raise HTTPException(status_code=500, detail=f"Could not process DOCX content: {e}") from e
    # Modify _read_pdf to accept bytes content
    def _read_pdf_content(self, pdf_content: bytes) -> str:
        text = []
        try:
            # pdfplumber can open file-like objects (BytesIO)
            with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(layout=True)
                    if page_text:
                        text.append(page_text)
        except Exception as e:
            # Handle pdfplumber errors, e.g., corrupted PDF
            logger.error(f"Error reading PDF content: {e}")
            raise HTTPException(status_code=500, detail=f"Could not process PDF content: {e}") from e
        return "\n".join(text)
    
    async def delete_blob_from_gcs(self, gcs_uri: str):
        try:
            credentials_path = "school-kg-7bd58d53b816.json"
            storage_client = storage.Client.from_service_account_json(credentials_path)
            
            blob = storage.Blob.from_string(gcs_uri, client=storage_client)

            # Удаление через run_in_executor чтобы не блочить event loop
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, blob.delete)
            
            logger.info(f"✅ Successfully deleted blob: {gcs_uri}")
            return JSONResponse(status_code=200, content={"message": "File successfully deleted."})

        except Exception as e:
            logger.error(f"💥 Error deleting blob {gcs_uri}: {e}")
            raise HTTPException(status_code=500, detail=f"Could not delete file from storage: {e}") from e
    async def delete_report_from_gcs(self, gcs_uri: str):
        try:
            
            credentials_path = "school-kg-7bd58d53b816.json"
            storage_client = storage.Client.from_service_account_json(credentials_path)
            
            blob = storage.Blob.from_string(gcs_uri, client=storage_client)

            # Удаление через run_in_executor чтобы не блочить event loop
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, blob.delete)
            
            logger.info(f"✅ Successfully deleted blob: {gcs_uri}")
            return JSONResponse(status_code=200, content={"message": "File successfully deleted."})

        except Exception as e:
            logger.error(f"💥 Error deleting blob {gcs_uri}: {e}")
            raise HTTPException(status_code=500, detail=f"Could not delete file from storage: {e}") from e

    # New function to download from GCS
    async def _download_gcs_blob(self, gcs_uri: str) -> bytes:
        try:
            credentials_path = "school-kg-7bd58d53b816.json"
            storage_client = storage.Client.from_service_account_json(credentials_path)
            # Use storage.Blob.from_string(gcs_uri, client=storage_client) for parsing gs:// URI
            blob = storage.Blob.from_string(gcs_uri, client=storage_client)

            # Download content as bytes asynchronously
            loop = asyncio.get_event_loop()
            content = await loop.run_in_executor(None, blob.download_as_bytes)
            return content
        except Exception as e:
            # Handle GCS errors (permissions, not found, etc.)
            logger.error(f"Error downloading from GCS {gcs_uri}: {e}")
            raise HTTPException(status_code=500, detail=f"Could not retrieve file from storage: {e}") from e

    # Update parse_pdf_to_text to handle GCS URI
    async def parse_pdf_to_text(self, gcs_uri: str) -> str:
        pdf_content = await self._download_gcs_blob(gcs_uri)
        loop = asyncio.get_event_loop()
        # Run the synchronous pdfplumber logic in an executor
        return await loop.run_in_executor(None, self._read_pdf_content, pdf_content)
    
    async def parse_docx_to_text(self, gcs_uri: str) -> str:
        docx_content = await self._download_gcs_blob(gcs_uri)
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._read_docx_content, docx_content)
    


    async def parse_docx(self, gcs_uri: str, vacancy_id: int) -> dict:
        if vacancy_id is None:
            raise HTTPException(status_code=400, detail="vacancy_id is required")

        text = await self.parse_docx_to_text(gcs_uri)

        stmt = (
            select(JobPosting)
            .where(JobPosting.id == vacancy_id)
            .options(selectinload(JobPosting.skills))
        )
        result = await self.db.execute(stmt)
        job = result.scalar_one_or_none()

        if not job:
            raise HTTPException(status_code=404, detail="Vacancy not found")

        skills = [SkillSchema(title=skill.title) for skill in job.skills]
        requirements = job.requirements

        return await analyze_resume(text, skills, requirements)
    
    async def parse_pdf(self, gcs_uri: str, vacancy_id: int) -> dict:
        if vacancy_id is None:
            raise HTTPException(status_code=400, detail="vacancy_id is required")

        text = await self.parse_pdf_to_text(gcs_uri)

        stmt = (
            select(JobPosting)
            .where(JobPosting.id == vacancy_id)
            .options(selectinload(JobPosting.skills))
        )
        result = await self.db.execute(stmt)
        job = result.scalar_one_or_none()

        if not job:
            raise HTTPException(status_code=404, detail="Vacancy not found")

        skills = [SkillSchema(title=skill.title) for skill in job.skills]
        requirements = job.requirements

        return await analyze_resume(text, skills, requirements)
